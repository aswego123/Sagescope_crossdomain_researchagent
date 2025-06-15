import streamlit as st
from web_search import WebSearch
from gemini_api import GeminiAPI
import time
from datetime import datetime
import docx
from docx.shared import Inches
import io
import os

# Set page config
st.set_page_config(
    page_title="SageScope - AI Research Agent",
    page_icon="🤖",
    layout="wide"
)

# Custom CSS
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
    }
    .agent-message {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .source-badge {
        background-color: #e0e0e0;
        padding: 0.2rem 0.5rem;
        border-radius: 0.3rem;
        font-size: 0.8rem;
        margin-left: 0.5rem;
    }
    .sidebar-section {
        margin: 1rem 0;
        padding: 1rem;
        background-color: #f0f2f6;
        border-radius: 0.5rem;
    }
    </style>
    """, unsafe_allow_html=True)

# Initialize session state
if 'search_results' not in st.session_state:
    st.session_state.search_results = None
if 'agent_thoughts' not in st.session_state:
    st.session_state.agent_thoughts = []

# Initialize components
web_search = WebSearch()
gemini_api = GeminiAPI()

def create_word_document(results, query):
    """Create a Word document with research results."""
    doc = docx.Document()
    
    # Add title
    doc.add_heading('Cross-Domain Research Documentation', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Research Query: {query}')
    doc.add_paragraph('---')
    
    # Add results by source
    sources = {}
    for result in results:
        source = result['source']
        if source not in sources:
            sources[source] = []
        sources[source].append(result)
    
    for source, source_results in sources.items():
        doc.add_heading(f'Results from {source.upper()}', level=1)
        
        for idx, result in enumerate(source_results, 1):
            doc.add_heading(f'{idx}. {result["title"]}', level=2)
            doc.add_paragraph(f'URL: {result["url"]}')
            doc.add_paragraph('Summary:')
            doc.add_paragraph(result['snippet'])
            doc.add_paragraph('---')
    
    return doc

def display_agent_thoughts(thoughts):
    """Display agent's thought process"""
    st.markdown("### 🤖 Agent's Analysis")
    for thought in thoughts:
        st.markdown(f"<div class='agent-message'>{thought}</div>", unsafe_allow_html=True)

def display_search_results(results):
    """Display search results in a formatted way"""
    st.markdown("### 📚 Search Results")
    
    for idx, result in enumerate(results, 1):
        with st.expander(f"{idx}. {result['title']}"):
            st.markdown(f"**Source:** {result['source'].upper()}")
            st.markdown(f"**URL:** [{result['url']}]({result['url']})")
            st.markdown("**Snippet:**")
            st.markdown(result['snippet'][:500] + "..." if len(result['snippet']) > 500 else result['snippet'])

def main():
    st.title("🤖 SageScope AI Research Agent")
    st.markdown("An autonomous research agent powered by Gemini Pro")
    
    # Enhanced Sidebar
    with st.sidebar:
        st.header("🤖 Agent Dashboard")
        
        # Research Settings
        st.markdown("### ⚙️ Research Settings")
        max_results = st.slider("Maximum results per source", 1, 10, 5)
        
        # About SageScope
        st.markdown("### ℹ️ About SageScope")
        st.markdown("""
        SageScope is an AI-powered research agent that:
        - Searches across multiple domains
        - Provides real-time insights
        - Generates comprehensive documentation
        - Adapts to your research needs
        """)
        
        # Agent Capabilities
        st.markdown("### 🛠️ Agent Capabilities")
        st.markdown("""
        The agent can:
        - Perform cross-domain research
        - Analyze in real-time
        - Verify source credibility
        - Summarize complex content
        - Generate detailed reports
        - Adapt search strategies
        - Process multiple sources
        - Extract key insights
        """)
        
        # Research Platforms
        st.markdown("### 🔍 Research Platforms")
        st.markdown("""
        Academic Sources:
        - arXiv: Latest research papers
        - Google Scholar: Academic citations
        - PubMed: Medical research
        
        Web Sources:
        - Tavily: Real-time web search
        - Wikipedia: General knowledge
        """)
        
        # Research Statistics
        if st.session_state.search_results:
            st.markdown("### 📊 Research Statistics")
            sources = {}
            for result in st.session_state.search_results:
                source = result['source']
                sources[source] = sources.get(source, 0) + 1
            
            for source, count in sources.items():
                st.markdown(f"**{source.upper()}:** {count} results")
    
    # Main content
    query = st.text_input("What would you like me to research?", placeholder="e.g., AI in Agriculture")
    
    if st.button("🤖 Start Research", type="primary"):
        if not query:
            st.error("Please enter a research query")
            return
            
        with st.spinner("Agent is researching..."):
            # Agent's initial thought
            st.session_state.agent_thoughts.append(f"🤔 Analyzing query: '{query}'")
            display_agent_thoughts(st.session_state.agent_thoughts)
            
            # Search across sources
            st.session_state.agent_thoughts.append("🔍 Searching across multiple sources...")
            display_agent_thoughts(st.session_state.agent_thoughts)
            
            results = web_search.search_all(query, max_results)
            st.session_state.search_results = results
            
            if not st.session_state.search_results:
                st.error("No results found. Please try a different query.")
                return
            
            # Agent's analysis
            st.session_state.agent_thoughts.append(f"📊 Found {len(results)} relevant sources")
            display_agent_thoughts(st.session_state.agent_thoughts)
            
            # Display search results
            display_search_results(st.session_state.search_results)
            
            # Create and save Word document
            doc = create_word_document(results, query)
            
            # Save to bytes
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            # Create download button
            st.download_button(
                label="📥 Download Research Documentation (Word)",
                data=docx_bytes,
                file_name="crossdomain_researchdocumentation.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Agent's final thoughts
            st.session_state.agent_thoughts.append("💡 Research complete. You can now explore the results above and download the documentation.")
            display_agent_thoughts(st.session_state.agent_thoughts)

if __name__ == "__main__":
    main() 